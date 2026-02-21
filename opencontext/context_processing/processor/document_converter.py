#!/usr/bin/env python
# -*- coding: utf-8 -*-

# Copyright (c) 2025 Beijing Volcano Engine Technology Co., Ltd.
# SPDX-License-Identifier: Apache-2.0

"""
Document Converter Helper Class

Provides document conversion and analysis functions:
- Document to images (PDF/DOCX/PPTX/images)
- Page-by-page analysis (PDF/DOCX): Extract text + detect visual elements
"""

import os
import tempfile
from pathlib import Path
from typing import List

from PIL import Image

from opencontext.utils.logging_utils import get_logger

logger = get_logger(__name__)


class PageInfo:
    """Page information container"""

    def __init__(
        self,
        page_number: int,
        text: str = "",
        has_visual_elements: bool = False,
        doc_images: List[Image.Image] = None,
    ):
        self.page_number = page_number
        self.text = text
        self.has_visual_elements = has_visual_elements  # Whether contains images/tables
        self.doc_images = doc_images or []  # Embedded images list (for DOCX only)

    def __repr__(self):
        return f"PageInfo(page={self.page_number}, text_len={len(self.text)}, visual={self.has_visual_elements}, images={len(self.doc_images)})"


class DocumentConverter:
    """Document Converter - read once, provide all information"""

    def __init__(self, dpi: int = 200):
        self.dpi = dpi

    def convert_to_images(self, file_path: str) -> List[Image.Image]:
        """Convert document to image list"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = Path(file_path).suffix.lower()
        logger.info(f"Converting document to images: {file_path} (type: {file_ext})")

        if file_ext == ".pdf":
            return self._convert_pdf_to_images(file_path)
        elif file_ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"]:
            return self._load_image(file_path)
        elif file_ext in [".pptx", ".ppt"]:
            return self._convert_pptx_to_images(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def _convert_pdf_to_images(self, pdf_path: str) -> List[Image.Image]:
        """Convert PDF to image list (using pypdfium2)"""
        try:
            import pypdfium2 as pdfium

            pdf = pdfium.PdfDocument(pdf_path)
            images = []
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                # Render page as PIL Image
                # scale parameter controls resolution: scale=1 corresponds to 72 DPI
                scale = self.dpi / 72.0
                pil_image = page.render(scale=scale).to_pil()
                if pil_image.mode != "RGB":
                    pil_image = pil_image.convert("RGB")
                images.append(pil_image)
            pdf.close()
            return images

        except Exception as e:
            logger.exception(f"Error converting PDF: {e}")
            raise

    def _load_image(self, image_path: str) -> List[Image.Image]:
        """Load single image"""
        logger.info(f"Loading image: {image_path}")
        try:
            img = Image.open(image_path)
            if img.mode != "RGB":
                img = img.convert("RGB")
            return [img]
        except Exception as e:
            logger.exception(f"Error loading image: {e}")
            raise

    def _convert_pptx_to_images(self, pptx_path: str) -> List[Image.Image]:
        """Convert PPTX to images (requires LibreOffice)"""
        logger.warning("PPTX processing requires LibreOffice")
        try:
            import subprocess

            with tempfile.TemporaryDirectory() as temp_dir:
                subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        temp_dir,
                        pptx_path,
                    ],
                    check=True,
                    capture_output=True,
                )

                pdf_filename = Path(pptx_path).stem + ".pdf"
                temp_pdf_path = os.path.join(temp_dir, pdf_filename)

                if not os.path.exists(temp_pdf_path):
                    raise FileNotFoundError(f"PDF conversion failed: {temp_pdf_path}")

                images = self._convert_pdf_to_images(temp_pdf_path)
            return images
        except subprocess.CalledProcessError as e:
            logger.error(
                "LibreOffice conversion failed. Install with: sudo apt-get install libreoffice"
            )
            raise RuntimeError(f"PPTX conversion failed: {e}")
        except Exception as e:
            logger.exception(f"Error converting PPTX: {e}")
            raise

    def analyze_pdf_pages(self, file_path: str, text_threshold: int = 50) -> List[PageInfo]:
        """
        Analyze each PDF page (one-time read, detect visual elements)
        """
        import pypdf

        page_infos = []

        with open(file_path, "rb") as pdf_file:
            pdf_reader = pypdf.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)

            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                # 2. Detect if has images/tables
                has_images = self._check_pdf_page_has_images(page)
                # 3. Determine if VLM is needed
                needs_vlm = has_images or len(text.strip()) < text_threshold

                page_info = PageInfo(
                    page_number=page_num + 1, text=text, has_visual_elements=needs_vlm
                )
                page_infos.append(page_info)
        return page_infos

    def _check_pdf_page_has_images(self, page) -> bool:
        """Check if PDF page contains images"""
        try:
            if "/Resources" not in page:
                return False
            resources = page["/Resources"]
            if "/XObject" not in resources:
                return False
            xobjects = resources["/XObject"].get_object()
            for obj_name in xobjects:
                xobject = xobjects[obj_name]
                if xobject.get("/Subtype") == "/Image":
                    return True
            return False
        except Exception:
            return False

    def analyze_docx_pages(self, file_path: str) -> List[PageInfo]:
        """
        Analyze DOCX (split by paragraphs, extract images for VLM)

        Strategy:
        1. Split paragraph groups by page break or fixed character count
        2. Convert tables to text, treat as normal text
        3. Extract images as PIL.Image objects, save to PageInfo.doc_images
        """
        import docx

        logger.info(f"Analyzing DOCX by paragraphs: {file_path}")
        try:
            doc = docx.Document(file_path)

            # 1. Split by page break or fixed character count
            page_groups = self._split_docx_into_groups(doc)

            if not page_groups:
                logger.warning(f"No content found in DOCX: {file_path}")
                return []

            # 2. Analyze each paragraph
            page_infos = []
            for page_num, group in enumerate(page_groups, start=1):
                text = group["text"]
                has_images = group["has_images"]
                doc_images = group.get("doc_images", [])

                needs_vlm = has_images
                page_info = PageInfo(
                    page_number=page_num,
                    text=text,
                    has_visual_elements=needs_vlm,
                    doc_images=doc_images,
                )
                page_infos.append(page_info)
            return page_infos

        except Exception as e:
            logger.exception(f"Error analyzing DOCX: {e}")
            raise

    def _split_docx_into_groups(self, doc, chars_per_group: int = 2000) -> list:
        """
        Split DOCX into paragraph groups (convert tables to text, extract images)
        """
        # 1. Build ordered mapping of document elements (paragraphs and tables)
        # Get true order of paragraphs and tables by parsing XML
        body_elements = self._get_body_elements(doc)

        groups = []
        current_paragraphs = []
        current_text_length = 0
        current_has_images = False
        current_doc_images = []

        # 2. Traverse document elements in order (paragraphs and tables)
        for element_type, element in body_elements:
            if element_type == "paragraph":
                paragraph = element
                # Check if has images
                para_images = self._extract_paragraph_images(paragraph, doc)
                if para_images:
                    current_has_images = True
                    current_doc_images.extend(para_images)

                has_page_break = self._has_page_break(paragraph)

                para_text = paragraph.text.strip()
                if para_text:
                    current_paragraphs.append(para_text)
                    current_text_length += len(para_text)

                should_split = has_page_break or current_text_length >= chars_per_group

                if should_split and current_paragraphs:
                    # Save current group
                    group_text = "\n\n".join(current_paragraphs)
                    groups.append(
                        {
                            "text": group_text,
                            "has_images": current_has_images,
                            "doc_images": current_doc_images,
                        }
                    )

                    current_paragraphs = []
                    current_text_length = 0
                    current_has_images = False
                    current_doc_images = []

            elif element_type == "table":
                table = element
                # Convert table to text
                table_text = self._table_to_text(table)
                if table_text:
                    current_paragraphs.append(f"\n=== Table ===\n{table_text}")
                    current_text_length += len(table_text)

        # 3. Save last group
        if current_paragraphs:
            group_text = "\n\n".join(current_paragraphs)
            groups.append(
                {
                    "text": group_text,
                    "has_images": current_has_images,
                    "doc_images": current_doc_images,
                }
            )

        # 4. If no groups, treat entire document as one group
        if not groups:
            all_text = "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            all_images = self._extract_all_images(doc)
            groups.append(
                {
                    "text": all_text,
                    "has_images": bool(all_images),
                    "doc_images": all_images,
                }
            )
        return groups

    def _get_body_elements(self, doc):
        """
        Get all elements (paragraphs and tables) in document body, return in order

        Returns: [('paragraph', paragraph_obj), ('table', table_obj), ...]
        """
        body_elements = []
        body = doc.element.body
        for child in body:
            if child.tag.endswith("p"):
                for paragraph in doc.paragraphs:
                    if paragraph._element == child:
                        body_elements.append(("paragraph", paragraph))
                        break
            elif child.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element == child:
                        body_elements.append(("table", table))
                        break
        return body_elements

    def _table_to_text(self, table) -> str:
        """
        Convert table to text format
        """
        table_rows = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):  # Only add non-empty rows
                table_rows.append(" | ".join(row_text))

        return "\n".join(table_rows)

    def _has_page_break(self, paragraph) -> bool:
        """Check if paragraph contains page break"""
        try:
            if hasattr(paragraph, "_element"):
                for run in paragraph.runs:
                    if hasattr(run, "_element"):
                        # Check <w:br w:type="page"/>
                        for child in run._element:
                            if child.tag.endswith("br"):
                                br_type = child.get(
                                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                                )
                                if br_type == "page":
                                    return True
            return False
        except Exception:
            return False

    def _extract_paragraph_images(self, paragraph, doc) -> List[Image.Image]:
        """
        Extract embedded images from paragraph

        Returns: [PIL.Image, ...]
        """
        images = []
        try:
            if not hasattr(paragraph, "_element"):
                return images

            # Traverse runs in paragraph
            for run in paragraph.runs:
                if not hasattr(run, "_element"):
                    continue
                # Find drawing elements
                drawing_elements = run._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                )
                for drawing in drawing_elements:
                    # Find blip elements in drawing (contains image references)
                    blip_elements = drawing.findall(
                        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                    )
                    for blip in blip_elements:
                        # Get image relationship ID
                        embed_attr = blip.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )
                        if embed_attr:
                            try:
                                # Get image data through relationship ID
                                image_part = doc.part.related_parts[embed_attr]
                                image_data = image_part.blob

                                # Convert image data to PIL.Image
                                import io

                                img = Image.open(io.BytesIO(image_data))
                                if img.mode != "RGB":
                                    img = img.convert("RGB")
                                images.append(img)
                                logger.debug(f"Extracted image from paragraph: {img.size}")

                            except Exception as e:
                                logger.warning(f"Failed to extract image: {e}")

        except Exception as e:
            logger.warning(f"Error extracting paragraph images: {e}")

        return images

    def _extract_all_images(self, doc) -> List[Image.Image]:
        """
        Extract all images from document

        Returns: [PIL.Image, ...]
        """
        images = []
        try:
            # Traverse all paragraphs
            for paragraph in doc.paragraphs:
                para_images = self._extract_paragraph_images(paragraph, doc)
                images.extend(para_images)

            logger.info(f"Extracted {len(images)} images from DOCX")
        except Exception as e:
            logger.warning(f"Error extracting all images: {e}")

        return images

    def analyze_markdown_pages(self, file_path: str, chars_per_group: int = 2000) -> List[PageInfo]:
        """
        Analyze Markdown file (group by heading level + character count, extract local images)

        Strategy:
        1. Group by # or ## headings (priority)
        2. If group exceeds chars_per_group, split by character count
        3. Extract local images for each group (![](path))
        4. Return PageInfo list (text + doc_images)
        """
        logger.info(f"Analyzing Markdown file: {file_path}")
        try:
            md_dir = Path(file_path).parent

            # Read Markdown file
            with open(file_path, "r", encoding="utf-8") as f:
                md_content = f.read()

            if not md_content.strip():
                logger.warning(f"Empty Markdown file: {file_path}")
                return []

            # Group by headings + character count
            groups = self._split_markdown_into_groups(md_content, md_dir, chars_per_group)

            if not groups:
                logger.warning(f"No content found in Markdown: {file_path}")
                return []

            # Build PageInfo list
            page_infos = []
            for page_num, group in enumerate(groups, start=1):
                text = group["text"]
                has_images = group["has_images"]
                doc_images = group.get("doc_images", [])

                page_info = PageInfo(
                    page_number=page_num,
                    text=text,
                    has_visual_elements=has_images,
                    doc_images=doc_images,
                )
                page_infos.append(page_info)

            return page_infos

        except Exception as e:
            logger.exception(f"Error analyzing Markdown: {e}")
            raise

    def _split_markdown_into_groups(
        self, md_content: str, md_dir: Path, chars_per_group: int = 2000
    ) -> list:
        """
        Group Markdown by headings + character count

        Strategy:
        1. Group by # or ## headings
        2. If a group exceeds chars_per_group, split it by character count
        3. Extract local images for each group
        """
        import re

        # Parse heading positions (# and ##)
        header_pattern = r"^(#{1,2})\s+(.+)$"
        lines = md_content.split("\n")

        groups = []
        current_lines = []
        current_text_length = 0

        for i, line in enumerate(lines):
            match = re.match(header_pattern, line)
            is_header = match is not None

            # Hit heading or reached character threshold
            should_split = (
                (is_header and current_lines)  # Hit new heading (and has content)
                or current_text_length >= chars_per_group  # Reached character threshold
            )

            if should_split:
                # Save current group
                group_text = "\n".join(current_lines)
                if group_text.strip():
                    doc_images, has_images = self._extract_markdown_images(group_text, md_dir)
                    groups.append(
                        {
                            "text": group_text,
                            "has_images": has_images,
                            "doc_images": doc_images,
                        }
                    )

                # Reset
                current_lines = []
                current_text_length = 0

            # Add current line
            current_lines.append(line)
            current_text_length += len(line)

        # Save last group
        if current_lines:
            group_text = "\n".join(current_lines)
            if group_text.strip():
                doc_images, has_images = self._extract_markdown_images(group_text, md_dir)
                groups.append(
                    {
                        "text": group_text,
                        "has_images": has_images,
                        "doc_images": doc_images,
                    }
                )

        return groups

    def _extract_markdown_images(self, md_text: str, md_dir: Path) -> tuple:
        """
        Extract local and remote images from Markdown text.

        Returns: (images: List[PIL.Image], has_images: bool)
        """
        import io
        import re
        import urllib.request

        images = []

        # Match ![alt](path) syntax
        pattern = r"!\[.*?\]\((.*?)\)"
        matches = re.findall(pattern, md_text)

        for img_path_str in matches:
            img_path_str = img_path_str.strip()

            try:
                if img_path_str.startswith(("http://", "https://")):
                    # Handle remote image by downloading it
                    with urllib.request.urlopen(img_path_str, timeout=10) as response:
                        image_data = response.read()
                        img = Image.open(io.BytesIO(image_data))
                        if img.mode != "RGB":
                            img = img.convert("RGB")
                        images.append(img)
                        logger.debug(
                            f"Successfully downloaded remote image: {img_path_str[:70]}..."
                        )

                elif not img_path_str.startswith("data:"):
                    # Handle local image
                    img_path = Path(img_path_str)

                    # Convert relative path to absolute path relative to the markdown file
                    if not img_path.is_absolute():
                        img_path = (md_dir / img_path).resolve()

                    if not img_path.exists():
                        logger.warning(f"Local image file not found: {img_path}")
                        continue

                    img = Image.open(img_path)
                    if img.mode != "RGB":
                        img = img.convert("RGB")
                    images.append(img)
                    logger.debug(f"Loaded local image: {img_path}")

            except Exception as e:
                logger.warning(f"Failed to load or download image '{img_path_str}': {e}")
                continue

        has_images = len(images) > 0
        if has_images:
            logger.info(f"Extracted {len(images)} images from Markdown text")

        return images, has_images
